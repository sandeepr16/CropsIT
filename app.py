from typing import OrderedDict
import docx
import jsonify
import requests
import pickle
import numpy as np
import sklearn
from sklearn.preprocessing import StandardScaler
import pyrebase
from flask import Flask, flash, redirect, render_template, request, session, abort, url_for
from twilio.rest import Client
import datetime
import time
from collections import OrderedDict
import pandas as pd

linkdf = pd.read_csv("Fertiliserslink.csv")
app = Flask(__name__)
model = pickle.load(open('croptype1.sav', 'rb'))
model2 = pickle.load(open('fertilizer1.sav', 'rb'))
# Add your own details
config = {
    "apiKey": "AIzaSyArgMlX6Bmor5u7R-kGhSXxhzwKNfMhjyY",
    "authDomain": "cropsit-a801f.firebaseapp.com",
    "databaseURL": "https://cropsit-a801f-default-rtdb.firebaseio.com/",
    "projectId": "cropsit-a801f",
    "storageBucket": "cropsit-a801f.appspot.com",
    "messagingSenderId": "52828763751",
    "appId": "1:52828763751:web:04fd7ff492d33fc89086c2",
    "measurementId": "G-5PRN42LB3F"
}

# initialize firebase
firebase = pyrebase.initialize_app(config)
auth = firebase.auth()
db = firebase.database()
person = {}
dataof = dict()

# Initialze person as dictionary
person = {"is_logged_in": False, "name": "",
          "email": "", "uid": "", "phoneno": ""}
# Login


@app.route('/')
def Home():
    return render_template('login.html')
# register


@app.route('/login')
def login():
    return render_template("login.html")


@app.route("/signup")
def signup():
    return render_template("signup.html")

#


dataof = {"crop": "", "fert": "", "date": "", "time": ""}


@app.route("/dashboard", methods=["GET"])
def dashboard():
    if person['is_logged_in'] == True:
        ls = []
        cr = ['Crop', 'Date', 'Fertiliser(NPK comp)', 'Time']
        ans = db.child("users").child(person["uid"]).child("dat").get()
        for user in ans.each():
            ls.append(user.val())
        print("11")
        print(ans.val())
        return render_template("dash.html", answer=ls, heading=cr)
    else:
        return redirect(url_for('login'))


@app.route('/home', methods=["GET"])
def home():
    if person["is_logged_in"] == True:
        data = db.child("users").get()
        na1 = data.val()[person["uid"]]["name"]
        print(na1)
        return render_template("index1.html", na=na1)
    else:
        return redirect(url_for('login'))


@app.route("/predict", methods=['POST', 'GET'])
def predict():
    if person["is_logged_in"] == True:
        if request.method == 'POST':
            number = int(request.form['number'])
            soiltype1 = int(request.form['soiltype'])
            o = list(map(int, str(number)))
            temp = int(str(o[0])+str(o[1]))
            humid = int(str(o[2])+str(o[3]))
            moist = int(str(o[4])+str(o[5]))
            soiltype = soiltype1
            nitro = int(str(o[6])+str(o[7]))
            potas = int(str(o[8])+str(o[9]))
            phos = int(str(o[10])+str(o[11]))
            fv = [temp, humid, moist, soiltype, nitro, potas, phos]
            fv = np.array(fv).reshape((1, -1))
            prediction = model.predict(fv)
            crop = {'Maize': 1, 'Sugarcane': 2, 'Cotton': 3, 'Tobacco': 4, 'Paddy': 5, 'Barley': 6,
                    'Wheat': 7, 'Millets': 8, 'Oilseeds': 9, 'Pulses': 10, 'GroundNuts': 11}
            crop_name = list(crop.keys())[list(
                crop.values()).index(prediction)]
            prediction1 = model2.predict(
                [[temp, humid, moist, soiltype, prediction, nitro, potas, phos]])
            fert_name = {'Urea': 1, 'DAP': 2, '14-35-14': 3,
                         '28-28': 4, '17-17-17': 5, '20-20': 6, '10-26-26': 7}
            y = list(fert_name.keys())[
                list(fert_name.values()).index(prediction1)]
            ts = time.time()
            st = datetime.datetime.fromtimestamp(
                ts).strftime('%d-%m-%Y')
            ti = datetime.datetime.fromtimestamp(ts).strftime('%H:%M:%S')

            dataof["crop"] = crop_name
            dataof["fert"] = y
            dataof['date'] = st
            dataof['time'] = ti
            print(dataof)
            db.child("users").child(person["uid"]).child("dat").push(dataof)
            ans = db.child("users").child(person["uid"]).child("dat").get()
            print(ans.val())
            account_sid = 'AC2749609bff78598eba194d52d5e517dd'
            auth_token = 'e80cd35cdf9ef191cd04b4237fd4f121'
            client = Client(account_sid, auth_token)
            data = db.child("users").get()
            ans = db.child("users").child("dat").get()
            person["num"] = data.val()[person["uid"]]["num"]
            person["name"] = data.val()[person["uid"]]["name"]
            na1 = data.val()[person["uid"]]["name"]
            sendernumber = ''
            sendernumber = str("+91")+str(person["num"])
            try:
                message = client.messages.create(
                    body='Hi' + str(person['name'])+' Cropname : ' +
                    str(crop_name)+' Fertilisername : ' + str(y),
                    from_='+14803861247',
                    to=sendernumber
                )
            except:
                pass
            # print(message.sid)
            doc = docx.Document("crops/"+str(crop_name)+".docx")
            all_paras = doc.paragraphs
            imgurl1 = "static/img/"+crop_name+".jpg"

            return render_template('output.html', crop=crop_name, fert=y, all=all_paras, imgurl2=imgurl1, na=na1)
        else:
            return render_template('index1.html')
    else:
        return redirect(url_for('login'))
# If someone clicks on login, they are redirected to /result


@app.route("/result", methods=["POST", "GET"])
def result():
    if request.method == "POST":  # Only if data has been posted
        result = request.form  # Get the data
        email = result["email"]
        password = result["pass"]
        print(email)
        try:
            # Try signing in the user with the given information
            user = auth.sign_in_with_email_and_password(email, password)
            # Insert the user data in the global person
            global person
            person["is_logged_in"] = True
            person["email"] = user["email"]
            person["uid"] = user["localId"]
            # Get the name of the user
            data = db.child("users").get()
            person["name"] = data.val()[person["uid"]]["name"]
            # Redirect to welcome page
            return redirect(url_for('home'))
        except:
            # If there is any error, redirect back to login

            return redirect(url_for('login'))
    else:
        if person["is_logged_in"] == True:
            return redirect(url_for('home'))
        else:
            return redirect(url_for('login'))

# If someone clicks on register, they are redirected to /register


@app.route("/register", methods=["POST", "GET"])
def register():
    if request.method == "POST":  # Only listen to POST
        result = request.form  # Get the data submitted
        email = result["email"]
        password = result["pass"]
        name = result["name"]
        number = result["num"]
        try:
            # Try creating the user account using the provided data
            auth.create_user_with_email_and_password(email, password)
            # Login the user
            user = auth.sign_in_with_email_and_password(
                email, password)
            # Add data to global person
            global person
            person["is_logged_in"] = True
            person["email"] = user["email"]
            person["uid"] = user["localId"]
            person["name"] = name
            person["num"] = number
            # Append data to the firebase realtime database
            data = {"name": name, "email": email, "num": number}
            db.child("users").child(person["uid"]).set(data)
            # Go to welcome page
            return redirect(url_for('home'))
        except:
            # If there is any error, redirect to register
            return redirect(url_for('signup'))

    else:
        if person["is_logged_in"] == True:
            return redirect(url_for('home'))
        else:
            return redirect(url_for('register'))


@app.route('/se', methods=['GET', 'POST'])
def se():
    return render_template("search.html")


@app.route('/search', methods=['GET', 'POST'])
def search():
    if request.method == "POST":
        s1 = request.form["cro"]
        all = searchresult(s1)
        imgurl = 'static/img/'+s1+'.jpg'
    return render_template("search1.html", sear=all, ss=imgurl)


def searchresult(s1):
    doc = docx.Document("crops/"+str(s1)+".docx")
    all = doc.paragraphs
    return all


@app.route('/fertiliser', methods=['GET', 'POST'])
def fertiliser():
    row_data = list(linkdf.values.tolist())
    print(row_data)
    return render_template("fertiliser.html", column_names=linkdf.columns.values, row_data=list(linkdf.values.tolist()),
                           zip=zip)


if __name__ == "__main__":
    app.run(debug=True)
